#!/usr/bin/env python3
"""render_docx.py — Allegro Launch Kit post-verify DOCX renderer.

ORDERING RULE: extract → generate → verify (on the .md) → render to .docx.
This script must only be called after verify.py exits 0 on the input file.

Design choice — --verified flag vs. internal verify call:
  Chosen: --verified flag.
  Reason: render_docx.py reads only the Markdown file; it has no natural need
  for specs.json. Requiring specs.json here solely to run the verifier would
  couple two unrelated concerns. The --verified flag is the pipeline's explicit
  assertion that verify.py ran clean — which makes the ordering rule visible
  in any shell script or Makefile that wires the steps together.

Usage:
    # 1. Verify first:
    python3 scripts/verify.py specs.json out_faq.md
    # 2. If exit 0, render:
    python3 scripts/render_docx.py out_faq.md out_faq.docx --verified
"""

import argparse
import re
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx",
          file=sys.stderr)
    sys.exit(2)


# ── Inline token patterns ─────────────────────────────────────────────────────
# [UNVERIFIED...] — single-line after pre_join_multiline()
_UNVERIFIED_RE = re.compile(r'\[UNVERIFIED[^\]]*\]', re.IGNORECASE)
# [datasheet p.N ...] or [p.N ...]
_CITATION_RE = re.compile(r'\[(?:datasheet\s+)?p\.\d+[^\]]*\]', re.IGNORECASE)
# **bold**
_BOLD_RE = re.compile(r'\*\*(.+?)\*\*')


# ── OOXML helpers ─────────────────────────────────────────────────────────────

def _highlight_yellow(run):
    """Apply Word yellow highlight to a run."""
    rPr = run._r.get_or_add_rPr()
    hl = OxmlElement('w:highlight')
    hl.set(qn('w:val'), 'yellow')
    rPr.append(hl)


def _para_shading(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def _para_border(para, color_hex, sides=None):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in (sides or ['top', 'left', 'bottom', 'right']):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color_hex)
        pBdr.append(el)
    pPr.append(pBdr)


def _add_hr(doc):
    """Horizontal rule via bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Inline formatter ──────────────────────────────────────────────────────────

def _add_runs(para, text):
    """Tokenize text and add styled runs to para.

    Priority: UNVERIFIED (yellow highlight + bold) → citation (grey 9pt italic)
              → bold (**...**) → plain text.
    Plain characters are buffered and flushed as a single run before each
    styled segment to keep the OOXML compact.
    """
    i = 0
    n = len(text)
    buf = []

    def flush():
        if buf:
            para.add_run(''.join(buf))
            buf.clear()

    while i < n:
        ch = text[i]

        if ch == '[':
            # UNVERIFIED?
            m = _UNVERIFIED_RE.match(text, i)
            if m:
                flush()
                run = para.add_run(m.group())
                run.bold = True
                _highlight_yellow(run)
                i = m.end()
                continue
            # Citation?
            m = _CITATION_RE.match(text, i)
            if m:
                flush()
                run = para.add_run(m.group())
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                i = m.end()
                continue

        if ch == '*' and i + 1 < n and text[i + 1] == '*':
            m = _BOLD_RE.match(text, i)
            if m:
                flush()
                run = para.add_run(m.group(1))
                run.bold = True
                i = m.end()
                continue

        buf.append(ch)
        i += 1

    flush()


# ── Multi-line UNVERIFIED normaliser ─────────────────────────────────────────

def _join_unverified(raw):
    """Collapse newlines inside [UNVERIFIED...] spans to spaces.

    UNVERIFIED blocks can span multiple lines in the Markdown source. Joining
    them here means the line-by-line parser always sees a single-line token,
    which _UNVERIFIED_RE (without DOTALL) then matches cleanly.
    """
    def _squash(m):
        return m.group().replace('\n', ' ').replace('  ', ' ')
    return re.sub(r'\[UNVERIFIED.*?\]', _squash,
                  raw, flags=re.DOTALL | re.IGNORECASE)


# ── Document structure helpers ────────────────────────────────────────────────

def _draft_banner(doc, filename):
    """Insert the DRAFT banner at the top of the document."""
    p = doc.add_paragraph()
    run = p.add_run(
        f"AUTO-GENERATED DRAFT  │  {date.today().isoformat()}  │  "
        f"Source: {filename}  │  "
        "Every spec is cited to the datasheet  │  "
        "[UNVERIFIED] items require human input before publication"
    )
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x7B, 0x34, 0x00)
    _para_shading(p, 'FFF3CD')
    _para_border(p, 'FFC107')
    doc.add_paragraph()


def _render_table(doc, rows):
    """Render a list of [col, ...] rows as a Word table."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell_text = row[ci] if ci < len(row) else ''
            cell = tbl.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            _add_runs(p, cell_text)
            if ri == 0:  # make header row bold
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


# ── Main converter ────────────────────────────────────────────────────────────

_ITALIC_HEADER = re.compile(r'^\*([^*].*[^*])\*$')
_BOLD_HEADING  = re.compile(r'^\*\*(.+)\*\*$')
_TABLE_SEP     = re.compile(r'^\|[\s\-:]+(?:\|[\s\-:]+)*\|?$')


def render_md_to_docx(md_path, output_path):
    """Convert a verified Allegro Markdown collateral file to a .docx."""
    raw = Path(md_path).read_text(encoding='utf-8')
    text = _join_unverified(raw)
    lines = text.split('\n')

    doc = Document()
    for section in doc.sections:
        section.left_margin  = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin   = Inches(1)
        section.bottom_margin = Inches(1)

    _draft_banner(doc, Path(md_path).name)

    i = 0
    pending_list = []   # [(text, is_checkbox), ...]

    def flush_list():
        for item_text, is_cb in pending_list:
            p = doc.add_paragraph(style='List Bullet')
            if is_cb:
                p.add_run('☐ ')  # ☐
            _add_runs(p, item_text)
        pending_list.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Blank line ────────────────────────────────────────────────────────
        if not stripped:
            flush_list()
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if stripped.startswith('# ') and not stripped.startswith('##'):
            flush_list()
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith('## '):
            flush_list()
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith('### '):
            flush_list()
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if stripped == '---':
            flush_list()
            _add_hr(doc)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if stripped.startswith('> '):
            flush_list()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── Markdown table ────────────────────────────────────────────────────
        if stripped.startswith('|'):
            flush_list()
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_s = lines[i].strip()
                if _TABLE_SEP.match(row_s):
                    i += 1
                    continue
                cols = [c.strip() for c in row_s.strip('|').split('|')]
                rows.append(cols)
                i += 1
            _render_table(doc, rows)
            continue

        # ── Italic section header line: *Text* (not **bold**) ─────────────────
        m_ih = _ITALIC_HEADER.match(stripped)
        if m_ih and not stripped.startswith('**'):
            flush_list()
            p = doc.add_paragraph()
            run = p.add_run(m_ih.group(1))
            run.italic = True
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            i += 1
            continue

        # ── List item: - [ ] ... or - ... ────────────────────────────────────
        if stripped.startswith('- '):
            rest = stripped[2:]
            is_cb = rest.startswith('[ ] ')
            if is_cb:
                rest = rest[4:]
            # Join indented continuation lines (non-UNVERIFIED continuation lines
            # that weren't collapsed by _join_unverified).
            while i + 1 < len(lines):
                nxt = lines[i + 1]
                nxt_s = nxt.strip()
                if (nxt.startswith('  ') and nxt_s
                        and not nxt_s.startswith(
                            ('- ', '#', '|', '> ', '---'))):
                    rest += ' ' + nxt_s
                    i += 1
                else:
                    break
            pending_list.append((rest, is_cb))
            i += 1
            continue

        # ── Standalone **Bold heading?** → H2 ────────────────────────────────
        m_bq = _BOLD_HEADING.match(stripped)
        if m_bq:
            flush_list()
            doc.add_heading(m_bq.group(1), level=2)
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────────────
        flush_list()
        p = doc.add_paragraph()
        _add_runs(p, stripped)
        i += 1

    flush_list()

    doc.save(output_path)
    print(f"  Rendered → {output_path}")


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(
        description=(
            "Render a verified Allegro Launch Kit .md file to a structured .docx.\n\n"
            "ORDERING RULE: run verify.py first; pass --verified only after exit 0.\n"
            "  python3 scripts/verify.py specs.json out_faq.md   # must exit 0\n"
            "  python3 scripts/render_docx.py out_faq.md out_faq.docx --verified"
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument('input_md',    help='Verified .md collateral file')
    ap.add_argument('output_docx', help='Output .docx path')
    ap.add_argument(
        '--verified',
        action='store_true',
        help=(
            "Assert that verify.py has been run on this file and exited 0. "
            "Required — render_docx.py refuses to render without this flag."
        ),
    )
    args = ap.parse_args()

    if not args.verified:
        print(
            "\n  ERROR: --verified flag not set.\n\n"
            "  render_docx.py may only run AFTER verify.py confirms all claims\n"
            "  trace to the datasheet. Run:\n\n"
            "    python3 scripts/verify.py specs.json <your_file.md>\n\n"
            "  If that exits 0, re-run this command with --verified.\n",
            file=sys.stderr,
        )
        sys.exit(1)

    md = Path(args.input_md)
    if not md.exists():
        print(f"ERROR: input file not found: {args.input_md}", file=sys.stderr)
        sys.exit(1)

    render_md_to_docx(str(md), args.output_docx)


if __name__ == '__main__':
    main()
